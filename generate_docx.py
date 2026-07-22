"""生成评分标准 docx 文档。"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()
style = doc.styles['Normal']
style.font.size = Pt(11)

# ── 标题 ──
title = doc.add_heading("CatoN News (CNN) 术力口周榜 — 评分标准", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("专注于 VOCALOID 新曲的 B站 周榜，每周二 12:00 截止统计。", style="Subtitle")


def make_table(headers: list[str], rows: list[list[str]]):
    """快捷创建格式化表格。"""
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers), style="Light Grid Accent 1")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for i, row in enumerate(rows, 1):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = val
    doc.add_paragraph()
    return table


# ══════════════════════════════════════
# 一、评分总公式
# ══════════════════════════════════════
doc.add_heading("一、评分总公式", level=1)

p = doc.add_paragraph()
r = p.add_run("S = (V + 25F + 2L + 8C) × exp(α × (1 − T/168)) × W")
r.bold = True
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

doc.add_paragraph()

make_table(
    ["符号", "含义", "取值"],
    [
        ["S", "最终得分", "—"],
        ["V", "播放量", "API 原始值"],
        ["F", "收藏数", "API 原始值"],
        ["L", "点赞数", "API 原始值"],
        ["C", "投币数", "API 原始值"],
        ["α", "时间衰减系数", "1.0"],
        ["T", "视频已发布小时数", "0 ~ 336"],
        ["168", "时间因子归一化基准（小时）", "固定值"],
        ["W", "投稿者权重", "1.00 / 1.01 / 1.03"],
        ["exp", "自然指数函数 e^x", "e 约等于 2.718"],
    ],
)

# ══════════════════════════════════════
# 二、基础分
# ══════════════════════════════════════
doc.add_heading("二、基础分", level=1)

p = doc.add_paragraph("基础分 = ")
p.add_run("播放 × 1 + 收藏 × 25 + 点赞 × 2 + 投币 × 8").bold = True

doc.add_paragraph("反映视频在 B站 的客观数据表现。")

make_table(
    ["数据项", "权重", "设计理由"],
    [
        ["播放 (V)", "×1", "传播广度 — 权重最低，易于刷量"],
        ["收藏 (F)", "×25", "作品质量 — 权重最高，用户主动保存行为"],
        ["点赞 (L)", "×2", "喜爱程度 — 轻量正向反馈"],
        ["投币 (C)", "×8", "深度认可 — 硬币有限，投币代表高度认可"],
    ],
)

# ══════════════════════════════════════
# 三、时间因子
# ══════════════════════════════════════
doc.add_heading("三、时间因子（后发补偿）", level=1)

p = doc.add_paragraph()
p.add_run("时间因子 = exp(α × (1 − T/168))").bold = True

doc.add_paragraph(
    "解决'先发优势'问题：周二截止时，上周三发的歌已有 6 天积累数据，"
    "而周二当天发的歌只有几小时数据。时间因子给后发新曲加权补偿。"
)

make_table(
    ["已发布时长 T", "时间因子", "备注"],
    [
        ["1 小时", "2.72", "刚发布的新曲，大幅补偿"],
        ["6 小时", "2.63", ""],
        ["12 小时", "2.50", "半天"],
        ["1 天 (24h)", "2.30", ""],
        ["2 天 (48h)", "2.00", ""],
        ["3 天 (72h)", "1.74", ""],
        ["5 天 (120h)", "1.35", ""],
        ["7 天 (168h)", "1.00", "满一周，不补偿"],
    ],
)

# ══════════════════════════════════════
# 四、投稿者权重
# ══════════════════════════════════════
doc.add_heading("四、投稿者权重", level=1)

p = doc.add_paragraph("人工审核时标记，权重差异很小，仅做微调。")

make_table(
    ["类型", "权重 (W)", "判定标准"],
    [
        ["本家 (original)", "1.03", "P主本人投稿 / 官方账号投稿"],
        ["授权 (authorized)", "1.01", "有 P主 / 官方授权证明的搬运"],
        ["未授权 (unauthorized)", "1.00", "默认值，无加成"],
    ],
)

# ══════════════════════════════════════
# 五、前置过滤
# ══════════════════════════════════════
doc.add_heading("五、前置过滤条件", level=1)

make_table(
    ["过滤项", "门槛", "不达标处理"],
    [
        ["最低播放量", ">= 100", "不进入排名计算"],
        ["最低投币数", ">= 20", "不进入排名计算"],
    ],
)

# ══════════════════════════════════════
# 六、术力口判定
# ══════════════════════════════════════
doc.add_heading("六、术力口判定（多维度打分）", level=1)

p = doc.add_paragraph()
p.add_run("判定公式: 强关键词分 + 标签分 + 歌姬分 - 排除分 - 时长扣分 >= 4 → 术曲").bold = True

doc.add_paragraph()

make_table(
    ["维度", "分值", "关键词示例"],
    [
        ["标题含强关键词", "+3", "VOCALOID / 术力口 / 术曲 / ボカロ / 初音 / ミク / SynthV / SV 等"],
        ["标签含术力口标签", "+5", "VOCALOID / 术力口 / ボカロ / UTAU / CeVIO / SynthV / NEUTRINO 等"],
        ["标题/标签含歌姬名", "+2", "初音ミク / 初音未来 / 重音テト / 可不 / 星界 / 洛天依 / 乐正绫 / 言和 等"],
        ["标题含排除词", "-3", "翻唱 / cover / reaction / 盘点 / 排行榜 / 教学 / 教程 / 非术曲"],
        ["时长 < 30 秒", "-5", "大概率是切片或短视频，非完整术曲作品"],
    ],
)

p = doc.add_paragraph()
p.add_run("注意：每类只取一次最高分，不重复计分。").italic = True

# ══════════════════════════════════════
# 七、完整计算示例
# ══════════════════════════════════════
doc.add_heading("七、完整计算示例", level=1)

doc.add_paragraph("假设一首周三发布的本家曲，截止时数据如下：")

example_data = [
    ("播放量 V", "50,000"),
    ("收藏数 F", "3,000"),
    ("点赞数 L", "8,000"),
    ("投币数 C", "1,500"),
    ("已发布小时数 T", "96 小时（4天）"),
    ("投稿者类型", "本家 (W = 1.03)"),
]
for label, val in example_data:
    doc.add_paragraph(f"  {label}: {val}", style="List Bullet")

doc.add_paragraph()

doc.add_paragraph(
    "基础分 = 50,000×1 + 3,000×25 + 8,000×2 + 1,500×8\n"
    "       = 50,000 + 75,000 + 16,000 + 12,000\n"
    "       = 153,000"
)

doc.add_paragraph(
    "时间因子 = exp(1.0 × (1 − 96/168))\n"
    "          = exp(0.4286)\n"
    "          = 1.54"
)

doc.add_paragraph(
    "最终得分 S = 153,000 × 1.54 × 1.03\n"
    "           = 242,689"
)

# ══════════════════════════════════════
# 八、收录范围
# ══════════════════════════════════════
doc.add_heading("八、收录范围与周期", level=1)

make_table(
    ["项目", "规则"],
    [
        ["统计周期", "两周（336小时）: 周二 12:00 ~ 两周后周二 12:00"],
        ["排名产出频率", "每周一次"],
        ["收录引擎", "VOCALOID / UTAU / Synthesizer V / CeVIO / NEUTRINO 等合成音声"],
        ["语言", "全语言（日/中/英/韩等）"],
        ["首次投稿要求", "必须是首次投稿，排除搬运换源/重复投稿"],
    ],
)

# ══════════════════════════════════════
# 九、总结
# ══════════════════════════════════════
doc.add_heading("九、一句话总结", level=1)

p = doc.add_paragraph()
r = p.add_run(
    "每首术曲的最终得分 = (播放×1 + 收藏×25 + 点赞×2 + 投币×8) "
    "× e^(1 − 发布天数/7) × 投稿者权重，"
    "得分最高的 10 首构成当周榜单。"
)
r.font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph("— CatoN News · 让每一首新术曲都有被看见的机会 —",
                  style="Intense Quote").alignment = WD_ALIGN_PARAGRAPH.CENTER

# 保存
doc.save("CatoN_News_评分标准.docx")
print("Done: CatoN_News_评分标准.docx")
